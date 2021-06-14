import os
import json
import uuid
import base64
from docx import Document
from celery import shared_task
from speech_analytics_platform.celery import app

from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (Mail, Attachment, FileContent, FileName,
                                    FileType, Disposition, ContentId)

from .services import json_to_diarized_txt_list

@app.task(bind=True, default_retry_delay=60, max_retries=5, acks_late=True)
def send_email(self, user_email_address, email_body, name, automatic_diarization):

    message = Mail(
        from_email=os.environ.get('EMAIL_USER'),
        to_emails=user_email_address,
        subject=name,
        html_content=email_body
    )

    try:
        print("Trying send email")
        sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
        response = sg.send(message)
        
    except Exception as e:
        self.retry(exc=e)
        
    print("Email sent!")

# from decoding_app.tasks import send_transcription_email
# send_transcription_email('flx.mcgregor@gmail.com', 'test', 'txt', '', 'name', False)
# send_transcription_email('flx.mcgregor@gmail.com', 'test', 'docx', '[{"startTime": 0.38, "endTime": 0.71, "word": "This", "confidence": 1.0, "speaker": "C"}, {"startTime": 0.71, "endTime": 0.89, "word": "is", "confidence": 1.0, "speaker": "C"}, {"startTime": 0.89, "endTime": 0.98, "word": "a", "confidence": 0.98, "speaker": "C"}, {"startTime": 0.98, "endTime": 1.52, "word": "test.", "confidence": 0.99, "speaker": "C"}, {"startTime": 3.06, "endTime": 3.51, "word": "This", "confidence": 1.0, "speaker": "B"}, {"startTime": 3.54, "endTime": 3.75, "word": "is", "confidence": 1.0, "speaker": "B"}, {"startTime": 3.75, "endTime": 3.87, "word": "a", "confidence": 0.86, "speaker": "B"}, {"startTime": 3.88, "endTime": 4.5, "word": "test.", "confidence": 0.99, "speaker": "B"}, {"startTime": 6.02, "endTime": 6.35, "word": "This", "confidence": 0.5, "speaker": "A"}, {"startTime": 6.38, "endTime": 6.59, "word": "is", "confidence": 0.92, "speaker": "A"}, {"startTime": 6.59, "endTime": 6.71, "word": "a", "confidence": 0.92, "speaker": "A"}, {"startTime": 6.71, "endTime": 7.31, "word": "test.", "confidence": 0.97, "speaker": "A"}]', 'name', True)
@app.task(bind=True, default_retry_delay=60, max_retries=5, acks_late=True)
def send_transcription_email(self, user_email_address, transcription, file_type, json_transcription, name, automatic_diarization):
    
    message = Mail(
        from_email=os.environ.get('EMAIL_USER'),
        to_emails=user_email_address,
        subject='Your transcription for ' + name,
        html_content='Find your transcription attached!'
    )
    
    attachment = Attachment()
    
    fn = os.path.splitext(name)[0]
    # choose file type to send transcription
    if file_type == 'txt':
        attachment.file_type = FileType('text/txt')
        attachment.file_name = FileName(f'{fn}.txt')
    elif file_type == 'docx':
        attachment.file_type = FileType('application/docx')
        attachment.file_name = FileName(f'{fn}.docx')
    else:
        # set transcript to json
        transcription = json_transcription
        attachment.file_type = FileType('application/json')
        attachment.file_name = FileName(f'{fn}.json')
       
    # Save the file to disk to encode for email
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    file_path = os.path.join(BASE_DIR, 'tmp')
    if not os.path.exists(file_path):
        os.makedirs(file_path)
    file_path = os.path.join(file_path, f'{str(uuid.uuid4().hex)}.tmp')

    if file_type == 'docx':
        document = Document()
        document.add_heading(fn, level=1)
        document.add_paragraph("")
        
        if automatic_diarization:
            transcription = json_to_diarized_txt_list(json.loads(json_transcription))
            for par in transcription:
                document.add_paragraph(par)
        else:
            document.add_paragraph(transcription)
            
        document.save(file_path)
        
    else:
        text_file = open(file_path, "w")
        text_file.write(transcription)
        text_file.close()
    with open(file_path, 'rb') as f:
        data = f.read()
        f.close()

    # cleanup
    os.remove(file_path)

    encoded = base64.b64encode(data).decode()

    # attach
    attachment.file_content = FileContent(encoded)
    attachment.disposition = Disposition('attachment')
    attachment.content_id = ContentId('Example Content ID')
    message.attachment = attachment

    
    try:
        print(f"Trying send for {name}")
        sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
        response = sg.send(message)
        print("Sent with attachment!")
    
    except Exception as e:
        print(e)
        self.retry(exc=e)
