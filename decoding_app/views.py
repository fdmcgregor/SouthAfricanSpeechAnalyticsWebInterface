import os
import json
import datetime
import requests

from django.shortcuts import render, get_object_or_404, redirect, reverse
from django.contrib.auth.models import User
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.views.generic import ListView, DetailView, CreateView, DeleteView
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponseRedirect
from django.utils.crypto import get_random_string

from django.utils.decorators import method_decorator

from .models import Post, SearchQuery, GlobalQuery
from .forms import SearchKeywordQueryForm, GlobalSearchForm, PostForm
from .services import get_account_info, search_transcript, get_similar_words, search_transcript_sequence, \
                    search_across_files, check_file_and_get_length, json_to_diarized_txt, json_to_diarized_txt_list
from .services import create_wordcloud
from .decode_service import decode_audio, save_api_response

max_word_len = 100
max_context_len = 10


class HomeView(ListView):
    model = Post
    template_name = 'decoding_app/home.html'


@method_decorator([login_required], name='dispatch')
class UserPostListView(LoginRequiredMixin, ListView):
    '''
    List view class which renders all posts by a user.
    '''

    model = Post
    template_name = 'decoding_app/user_posts.html'
    context_object_name = 'posts'
    ordering = ['-date_posted']
    paginate_by = 5

    def get_queryset(self):
        user = get_object_or_404(User, username=self.kwargs.get('username'))
        posts = Post.objects.filter(author=user)\
                            .filter(delete_after_transcription=False)\
                            .order_by('-date_posted')
        return posts


@method_decorator([login_required], name='dispatch')
class PostDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    '''
    DetailView class which displays details of a specific audio file showing
    its transcription and allows the user to interact with the it.
    '''

    model = Post

    def get_context_data(self, **kwargs):
        '''
        Called apon loading the page. It will render the trascription, wordcloud and
        old search queries for a particular post.
        '''
        context = super(PostDetailView, self).get_context_data(**kwargs)

        
        # word stats
        if len(context['post'].words_json) > 0:
            context['word_stats'] = json.loads(context['post'].words_json)
            if context['post'].done_decoding:
                if len(context['post'].transcription) > 0:
                    context['wordcloud'] = create_wordcloud(context['post'].transcription, context['post'].language)

        # split transcription for readmore or less
        context['post'].transcription_less = ' '.join(context['post'].transcription.split()[:max_word_len])
        context['post'].transcription_more = ' '.join(context['post'].transcription.split()[max_word_len:])

        # add a form for search queries
        context['form'] = SearchKeywordQueryForm

        # load previous search queries
        context['searchqueries'] = SearchQuery.objects.filter(post_pk=context['post'].pk).order_by('-date_searched')[0:5]

        # query specific logic
        if len(context['searchqueries']) > 0:

            keyword = str(context['searchqueries'][0])
            context['keyword'] = keyword

            # search transcript
            if len(keyword.split()) == 1:
                res_c, res_l, res_r = search_transcript(keyword, context['post'].transcription, \
                                                json.loads(context['post'].words_json), max_context_len)
            else: 
                res_c, res_l, res_r = search_transcript_sequence(keyword, context['post'].transcription, \
                                                                json.loads(context['post'].words_json), max_context_len)
                
            context['searchresult_stats_c'] = json.dumps(res_c)
            context['searchresult_stats_l'] = json.dumps(res_l)
            context['searchresult_stats_r'] = json.dumps(res_r)
            context['num_occurences'] = len(res_c)

        return context

    def post(self, request, slug):

        if request.is_ajax():

            keyword = str(request.POST.get('search_query')).lower()
            
            # find similar    
            if request.POST.get('find_similar') == 'true':

                lang = Post.objects.filter(slug=slug)[0].language
                if len(keyword.split()) == 1:
                    similar_words = get_similar_words(keyword, lang)
                else: 
                    # just take the first word for suggestions
                    similar_words = get_similar_words(keyword.split()[0], lang)

                context = {'suggestions': similar_words}
                return render(request, 'decoding_app/similar_words.html', context)

            # search results and save
            elif request.POST.get('search_results') == 'true':

                # save query
                post = Post.objects.filter(slug=slug)[0]
                SearchQuery.objects.create(keyword=keyword, post_pk=post.pk, author=request.user)

                # search transcript
                if len(keyword.split()) == 1:
                    res_c, res_l, res_r = search_transcript(keyword, post.transcription, \
                                                    json.loads(post.words_json), max_context_len)
                else: 
                    res_c, res_l, res_r = search_transcript_sequence(keyword, post.transcription, \
                                                                    json.loads(post.words_json), max_context_len)
                
                context = {}
                context['searchresult_stats_c'] = json.dumps(res_c)
                context['searchresult_stats_l'] = json.dumps(res_l)
                context['searchresult_stats_r'] = json.dumps(res_r)
                context['num_occurences'] = len(res_c)
                                
                context['keyword'] = keyword
                
                return render(request, 'decoding_app/search_results.html', context)

                    
            # render search history div
            else:
                context = {}
                post = Post.objects.filter(slug=slug)[0]
                context['searchqueries'] = SearchQuery.objects.filter(post_pk=post.pk).order_by('-date_searched')[0:5]
                return render(request, 'decoding_app/last_searched.html', context)

        # post with django if coming from global search
        form = SearchKeywordQueryForm(request.POST)
        if form.is_valid():
            # save keyword
            obj = form.save(commit=False)
            obj.author = request.user
            obj.post_pk = pk
            obj.save()

            return HttpResponseRedirect(reverse('post-detail', kwargs={'slug': slug}))

    def test_func(self):
        post = self.get_object()
        if self.request.user == post.author:
            return True
        return False

    def form_valid(self, form):
        form.instance.author = self.request.user
        return super().form_valid(form)

@login_required
def post_create_view(request):
    
    init={'delete_after_transcription': request.user.profile.delete_after_transcription,\
                             'sampling_freq': request.user.profile.sampling_freq,
                             'language': request.user.profile.language,
                             'file_type': request.user.profile.file_type,
                             'automatic_punctuation': request.user.profile.automatic_punctuation,
                             'automatic_diarization': request.user.profile.automatic_diarization,
                            }
    context = {}
    form = PostForm(initial=init)
    context['form'] = form 
    
    return render(request, "decoding_app/post_form.html", context) 


class PostDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Post
    success_url = '/'

    def test_func(self):
        post = self.get_object()
        if self.request.user == post.author:
            return True
        return False

    def delete(self, request, *args, **kwargs):
        """
        Call the delete() method on the fetched object and then redirect to the
        success URL.
        """
        self.object = self.get_object()
        success_url = self.get_success_url()
        
        post = self.object
        
        dt = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        print(f"Info {dt}: Delete post user:{request.user} post:{post}")
        
        post.delete_after_transcription = True
        post.audio_file.delete()
        post.transcription = "Not saved"
        post.words_json = "Not saved"
        post.save()

        return redirect('home')

from django.views.generic import View
from django.http import HttpResponse
import csv
from docx import Document

class FileView(LoginRequiredMixin, View):
    '''
    View to serve files to download
    '''
    def get(self, request, *args, **kwargs):
        
        #post = Post.objects.filter(pk=kwargs['pk'])[0]
        post = Post.objects.filter(slug=kwargs['slug'])[0]
        # choose extension
        file_type = kwargs['file_type']
        # get the post name 
        fn = os.path.splitext(post.title)[0]
        cd = f'attachment; filename="{fn}.{file_type}"'

        # mangage diarization
        if post.automatic_diarization:
            transcript = json_to_diarized_txt(json.loads(post.words_json))
        else:
            transcript = post.transcription
        
        if file_type == 'txt':
            response = HttpResponse(transcript, content_type='application/text')
        elif file_type == 'docx':
            
            document = Document()
            document.add_heading(fn, level=1)
            document.add_paragraph("")
            
            if post.automatic_diarization:
                transcript = json_to_diarized_txt_list(json.loads(post.words_json))
                for par in transcript:
                    document.add_paragraph(par)
            else:
                document.add_paragraph(transcript)
                
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            document.save(response)

        else:
            response = HttpResponse(post.words_json, content_type='application/json')
        
        response['Content-Disposition'] = cd
            
        return response

@login_required
def global_search(request):
    '''
    View that renders a form for the serach across files
    '''
    context = {}
    context['num_posts'] = Post.objects.filter(author=request.user).count()
        
    if request.method == 'POST':
        
        keyword = str(request.POST.get('keyword')).lower()
        
        # (1) redirect and search specific post
        if request.POST.get('global_search_redirect') == 'true':    
            
            keyword = str(request.POST.get('keyword')).lower()
            
            pk = str(request.POST.get('post_id'))
            post = Post.objects.filter(pk=pk)[0]
            
            # create query before page loads
            SearchQuery.objects.create(keyword=keyword, post_pk=post.pk, author=request.user)
            return HttpResponseRedirect(reverse('post-detail', kwargs={'slug': post.slug}))
          
        # (2) perform search on given keyword
        else:        
            # do global search
            form = GlobalSearchForm(request.POST)
            if form.is_valid():
                # save query
                obj = form.save(commit=False)
                obj.author = request.user
                obj.save()

                # get transcripts which contain keyword
                w = obj.keyword.lower()
                res_counts = search_across_files(w, request.user)
                context['results_keyword'] = w
                context['results'] = res_counts
            

    else:
        form = GlobalSearchForm()


    # load previous search queries and suggestions
    context['searchqueries'] = GlobalQuery.objects.filter(author=request.user).order_by('-date_searched')[0:5]
    
    if len(context['searchqueries']) > 0:
        keyword = str(context['searchqueries'][0])
        context['keyword'] = keyword
        
        # assume language is multi
        lang = 'mul'
        if len(keyword.split()) == 1:
            context['suggestions'] = get_similar_words(keyword, lang)
        else: 
            context['suggestions'] = get_similar_words(keyword.split()[0], lang)
        
    context['form'] = form

    return render(request, 'decoding_app/global_search.html', context)



def about(request):
    context = {'title': 'About'}
    return render(request, 'decoding_app/about.html', context)

from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse


def verify_and_upload(request):
    
    '''
    Before posting check that all files are valid and together it does not exceed limit.
    Support for uploading any media file (even video), support for uploading multiple files
    which we handle by sending individual requests once validated.
    '''
    
    if request.method == 'POST':
        
        files = [request.FILES[key] for key in request.FILES] 
        numspeakers_select = request.POST.get('numspeakers_select', False)

        # check each file
        total_duration = 0
        file_lengths = []
        file_num_channels = []
        fnames = []
        fnames_response = ''
        # ----------------------------------------
        # (1) validation step
        # ----------------------------------------
        for f in files: 
            

            f_name = f.name
            media, dur, num_channels = check_file_and_get_length(f)
            
            fnames_response += f'{f_name}, ' 
            fnames.append(f_name)

            dt = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            print(f"Info {dt}: File is_media:{media} duration:{dur}")
            file_lengths.append(dur)
            file_num_channels.append(num_channels)
            total_duration += dur
            
            # redirect if any file is not media
            if media != 1:
                messages.warning(request, f'Problem with file "{f_name}". Are you sure file contains audio?')
                return JsonResponse({"message": "not media"})
            
        # check that the account has not exceeded free trial or not suspended
        if request.user.profile.total_seconds + total_duration > int(os.environ.get('FREE_TRIAL_SECONDS')):
            if request.user.profile.allow_decoding_overide != True:
                
                remain = int(os.environ.get('FREE_TRIAL_SECONDS')) - request.user.profile.total_seconds 
                messages.warning(request, f'Upload exceeds trial. You have {remain}s left. Contact us to continue using this service.')
                return JsonResponse({"message": "expired"})

        # ----------------------------------------
        # (2) upload
        # ----------------------------------------
        
        messages.success(request, f"Busy transcribing {fnames_response[0:-2]}! \
                 You will get an email with the transcription or view it in your uploads.")

        # get metadata with post form
        form = PostForm(request.POST)
        
        for i, f in enumerate(files):    
       
            # (2.1) Save entry
            if form.is_valid():
                obj = form.save(commit=False)

            # if NB and stereo assume one speaker per channel
            if obj.sampling_freq == 'NB' and file_num_channels[i] == 2:
                obj.automatic_diarization = True

            obj.title = fnames[i]
            obj.author = request.user
            obj.audio_file = f
                
            obj.audio_length = file_lengths[i]
            if obj.automatic_diarization:
                if numspeakers_select:
                    obj.num_speakers = numspeakers_select
                
            obj.audio_length = file_lengths[i]
            obj.slug = get_random_string(length=32)
            obj.save()
                
            # (2.2) begin decoding (use .delay to use task runner to manage async process)
            print(f"Info {dt}: Starting job:{obj.pk} {obj.slug}")
            decode_audio.delay(obj.pk)
                    
                    
        return JsonResponse({"message": f"{fnames_response[0:-2]} uploaded!"})



def contact(request):
    context = {'title': 'Contact Us',
               'email': 'info@saigen.co.za',
               'tel': '+27 (84) 951 9002',
               'address': ['35 Brickfield Rd', 'Brickfield Canvas', 'Salt River', 'Cape Town 7925'],
               'web_address': 'https://saigen.co.za'
               }

    return render(request, 'decoding_app/contact_us.html', context)


@login_required
def account(request):

    # get account info for current user
    num_files, num_seconds = get_account_info(request.user)

    context = {'title': 'Account Information',
               'numfiles': num_files,
               'numseconds': num_seconds,
               }

    return render(request, 'decoding_app/account_info.html', context)


def expired(request):
    context = {'title': 'Your free trial is up or your account is outstanding.'}
    context['body'] = "To continue using our service please contact us"
    return render(request, 'decoding_app/expired.html', context)


from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
@csrf_exempt
def jobcomplete(request):

    dt = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"Info {dt}: Job complete {request.headers}")

    if request.method == 'GET':
        return HttpResponse('<h1>Return jobs here</h1>')

    elif request.method == 'POST':
        
        # validate
        # ----------------
        auth = request.headers.get("X_API_KEY", "")
        if auth != os.getenv('LAMBDA_API_KEY'):
            print(f"Info {dt}: Unauthorized ({auth})")
            return JsonResponse({'result': 'Error','message': 'Unauthorized'})

        user_id = request.headers.get("USER_ID", "")
        if len(user_id) > 0:
            if not User.objects.filter(pk=user_id).exists():
                print(f"Info {dt}: Unkown user ({user_id})")
                return JsonResponse({'result': 'Error','message': 'Unkown user'})
        else:
            print(f"Info {dt}: No user id")
            return JsonResponse({'result': 'Error', 'message': 'Specify USER_ID'})

        job_id = request.headers.get("JOB_ID", "")
        if len(job_id) > 0:
            if not Post.objects.filter(slug=job_id).exists():
                print(f"Info {dt}: Unkown job ({job_id})")
                return JsonResponse({'result': 'Error','message': 'Unkown job'})
        else:
            print(f"Info {dt}: No job id")
            return JsonResponse({'result': 'Error', 'message': 'Specify JOB_ID'})

        if Post.objects.filter(slug=job_id)[0].author != User.objects.filter(pk=user_id)[0]:
            print(f"Info {dt}: Job ({job_id}) does not belong to user ({user_id})")
            return JsonResponse({'result': 'Error','message': 'Incorrect job or user'})
        
        status = request.headers.get("STATUS", "")
        if status not in ['ERROR', 'SUCCESS']:
            print(f"Info {dt}: Incorrect status")
            return JsonResponse({'result': 'Error','message': 'Incorrect STATUS - Please return (ERROR | SUCCESS)'})

        # save
        # ----------
        print(f"Info {dt}: Saving response")
        s = request.FILES["file"].read()
        json_payload=json.loads(s)
        save_api_response(job_id, json_payload)

        return JsonResponse({'result': 'Success','message': 'Successfully recieved and recorded'}, status=200)